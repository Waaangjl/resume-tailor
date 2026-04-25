"""Convert a Word resume (.docx) to LaTeX using python-docx + LLM reconstruction."""

from pathlib import Path

import llm
from build import strip_tex_fence
from prompts import WORD_TO_LATEX_PROMPT

_DEFAULT_TEMPLATE = Path(__file__).parent / "resumes" / "sample_resume.tex"


def _extract_text(docx_path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx not installed — run: pip install python-docx"
        )
    doc = Document(str(docx_path))
    seen: set[str] = set()
    lines: list[str] = []

    def _add(text: str) -> None:
        t = text.strip()
        if t and t not in seen:
            seen.add(t)
            lines.append(t)

    for para in doc.paragraphs:
        _add(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(para.text)

    return "\n".join(lines)


def convert(docx_path: Path, model: str, template_path: Path | None = None) -> str:
    """Extract text from a .docx resume and reconstruct it as LaTeX.

    Returns the LaTeX source as a string. Caller is responsible for saving it.
    """
    template = (template_path or _DEFAULT_TEMPLATE).read_text(encoding="utf-8")
    text = _extract_text(docx_path)
    if not text.strip():
        raise ValueError(f"No text found in {docx_path.name} — is it a scanned PDF saved as .docx?")
    raw = llm.call(
        WORD_TO_LATEX_PROMPT.format(resume_text=text, template=template),
        model,
        timeout=180,
    )
    return strip_tex_fence(raw)
